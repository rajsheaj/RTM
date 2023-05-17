from docx.api import Document
import re
import os
from os import listdir
from os.path import isfile, join
import pandas as pd

class ExtractTestProtocols:

    def __init__(self, path):
        # self.files_path = [os.path.abspath(x) for x in os.listdir(path)]
        self.files_path = [os.path.join(path, f) for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
        # self.files_path1 = [os.path.join(path, f) for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
        print(self.files_path)
        self.df1 = pd.DataFrame(columns = ['Test_Step_ID', 'Test_ID','Protocol_Name','Req_ID'])
        self.df2 = pd.DataFrame(columns=['Protocol_Name','Test_Name', 'Test_ID'])

    def extract_test_steps(self):

        self.extract_test_name()

        for file in self.files_path:
            head, tail = os.path.split(file)
            print(tail)
            doc = Document(file)
            for table in doc.tables:
                if len(table.columns) == 6:

                    for row in table.rows[1:]:
                        test_step_id = row.cells[0].text
                        test_id = test_step_id[:test_step_id.rfind('.')]
                        test_step_req_id = row.cells[5].text
                        test_step_req_id = re.sub('[^A-Za-z0-9_]+', '', test_step_req_id)
                        newRow = {"Test_Step_ID": test_step_id, "Test_ID": test_id,"Protocol_Name": tail,"Req_ID": test_step_req_id}
                        row_df = pd.DataFrame([newRow])
                        self.df1 = pd.concat([self.df1, row_df], ignore_index=True)
                        self.df1 = (self.df1
                                .assign(Req_ID=self.df1['Req_ID'].str.split(r'(?<=.)(?=SRS)'))
                                .explode('Req_ID')
                               )

                        temp_df = self.df1.merge(self.df2, on=['Test_ID', 'Protocol_Name'], how='left')


        return temp_df

    def extract_test_name(self):
        for file in self.files_path:
            head, tail = os.path.split(file)
            #print(tail)
            doc = Document(file)

            for paragraph in doc.paragraphs:
                if paragraph.style.name == 'Heading 2':
                    if paragraph.text.find("Test Case ID / Title:") != -1:
                        ind = paragraph.text.find("Test Case ID / Title:")
                        testname = paragraph.text[(21 + ind):]
                        testid = paragraph.text[:ind].strip()
                        newRow = {"Protocol_Name": tail, "Test_Name": testname, "Test_ID": testid}
                        row_df = pd.DataFrame([newRow])
                        self.df2 = pd.concat([self.df2, row_df], ignore_index=True)

        # return self.df2

